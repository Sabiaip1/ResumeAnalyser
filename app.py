import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Р¤СѓРЅРєС†РёСЏ РґР»СЏ РёР·РІР»РµС‡РµРЅРёСЏ С‚РµРєСЃС‚Р° РёР· PDF
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# Р¤СѓРЅРєС†РёСЏ РґР»СЏ РёР·РІР»РµС‡РµРЅРёСЏ С‚РµРєСЃС‚Р° РёР· DOCX
def extract_text_from_docx(file):
    doc = Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# Р¤СѓРЅРєС†РёСЏ РґР»СЏ РѕРїСЂРµРґРµР»РµРЅРёСЏ РєР°С‚РµРіРѕСЂРёРё СЂРµР·СЋРјРµ
def determine_resume_category(text):
    if "СЃС‚СѓРґРµРЅС‚" in text.lower() or "СѓС‡С‘Р±Р°" in text.lower():
        return "student"
    elif "РѕРїС‹С‚ СЂР°Р±РѕС‚С‹" in text.lower() or "РїСЂРѕС„РµСЃСЃРёРѕРЅР°Р»" in text.lower():
        return "professional"
    return "general"

# Р—Р°РіР»СѓС€РєР° РґР»СЏ СЂРµРєРѕРјРµРЅРґР°С†РёР№
def generate_recommendations(category):
    if category == "student":
        return [
            "Р”РѕР±Р°РІСЊС‚Рµ РїРѕСЂС‚С„РѕР»РёРѕ РІР°С€РёС… СѓС‡РµР±РЅС‹С… РїСЂРѕРµРєС‚РѕРІ.",
            "РЈРєР°Р¶РёС‚Рµ РєСѓСЂСЃС‹ Рё СЃРµСЂС‚РёС„РёРєР°С‚С‹, РєРѕС‚РѕСЂС‹Рµ РІС‹ РїСЂРѕС€Р»Рё.",
            "РЎС„РѕРєСѓСЃРёСЂСѓР№С‚РµСЃСЊ РЅР° РЅР°РІС‹РєР°С… Рё РїРѕС‚РµРЅС†РёР°Р»Рµ."
        ]
    elif category == "professional":
        return [
            "РЈРєР°Р¶РёС‚Рµ РєРѕРЅРєСЂРµС‚РЅС‹Рµ РґРѕСЃС‚РёР¶РµРЅРёСЏ РІ С†РёС„СЂР°С… (РЅР°РїСЂРёРјРµСЂ, СѓРІРµР»РёС‡РёР» РїСЂРѕРґР°Р¶Рё РЅР° 20%).",
            "Р”РѕР±Р°РІСЊС‚Рµ СЃСЃС‹Р»РєРё РЅР° РІР°С€Рё СЂР°Р±РѕС‚С‹ РёР»Рё РїСЂРѕРµРєС‚С‹.",
            "РЎРґРµР»Р°Р№С‚Рµ Р°РєС†РµРЅС‚ РЅР° РєР»СЋС‡РµРІС‹С… РЅР°РІС‹РєР°С…, СЃРІСЏР·Р°РЅРЅС‹С… СЃ РІР°РєР°РЅСЃРёРµР№."
        ]
    else:
        return [
            "РЎРѕРєСЂР°С‚РёС‚Рµ РѕР±СЉС‘Рј СЂРµР·СЋРјРµ РґРѕ 1-2 СЃС‚СЂР°РЅРёС†.",
            "РџСЂРѕРІРµСЂСЊС‚Рµ С‚РµРєСЃС‚ РЅР° РЅР°Р»РёС‡РёРµ РѕСЂС„РѕРіСЂР°С„РёС‡РµСЃРєРёС… РѕС€РёР±РѕРє.",
            "Р”РѕР±Р°РІСЊС‚Рµ РєРѕРЅС‚Р°РєС‚РЅС‹Рµ РґР°РЅРЅС‹Рµ, РµСЃР»Рё РёС… РЅРµС‚."
        ]

# Р¤СѓРЅРєС†РёСЏ РґР»СЏ СЃРѕР·РґР°РЅРёСЏ PDF
def generate_pdf(recommendations):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    c.setFont("Helvetica", 12)
    c.drawString(100, 750, "Р РµРєРѕРјРµРЅРґР°С†РёРё РїРѕ СѓР»СѓС‡С€РµРЅРёСЋ СЂРµР·СЋРјРµ:")

    y_position = 730
    for rec in recommendations:
        c.drawString(100, y_position, f"- {rec}")
        y_position -= 20

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer

# Р¤СѓРЅРєС†РёСЏ РґР»СЏ СЃРѕР·РґР°РЅРёСЏ DOCX
def generate_docx(recommendations):
    doc = Document()
    doc.add_heading('Р РµРєРѕРјРµРЅРґР°С†РёРё РїРѕ СѓР»СѓС‡С€РµРЅРёСЋ СЂРµР·СЋРјРµ', 0)

    for rec in recommendations:
        doc.add_paragraph(f"- {rec}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# РџСЂРѕРіСЂРµСЃСЃ-Р±Р°СЂ РґР»СЏ РѕР±СЂР°Р±РѕС‚РєРё С„Р°Р№Р»Р°
def file_processing_progress_bar():
    progress_bar = st.progress(0)
    for i in range(100):
        progress_bar.progress(i + 1)

# РРЅС‚РµСЂР°РєС‚РёРІРЅС‹Рµ РЅР°СЃС‚СЂРѕР№РєРё РґР»СЏ СЂРµРєРѕРјРµРЅРґР°С†РёР№
def interactive_recommendations():
    st.sidebar.title("РќР°СЃС‚СЂРѕР№РєР° СЂРµРєРѕРјРµРЅРґР°С†РёР№")
    recommendation_type = st.sidebar.radio("Р’С‹Р±РµСЂРёС‚Рµ С‚РёРї СЂРµРєРѕРјРµРЅРґР°С†РёР№", ["Р”Р»СЏ СЃС‚СѓРґРµРЅС‚РѕРІ", "Р”Р»СЏ РїСЂРѕС„РµСЃСЃРёРѕРЅР°Р»РѕРІ", "РћР±С‰РёРµ"])
    return recommendation_type

# Р—Р°РіРѕР»РѕРІРѕРє РїСЂРёР»РѕР¶РµРЅРёСЏ
st.title("РђРЅР°Р»РёР·Р°С‚РѕСЂ Р РµР·СЋРјРµ РЅР° РѕСЃРЅРѕРІРµ РќРµР№СЂРѕСЃРµС‚Рё", anchor="top")

# Р‘РѕРєРѕРІР°СЏ РїР°РЅРµР»СЊ
st.sidebar.title("РќР°СЃС‚СЂРѕР№РєРё")
st.sidebar.write("Р’С‹Р±РµСЂРёС‚Рµ РїР°СЂР°РјРµС‚СЂС‹ РґР»СЏ Р°РЅР°Р»РёР·Р°:")
color_coding = st.sidebar.checkbox("Р¦РІРµС‚РѕРІР°СЏ РєРѕРґРёСЂРѕРІРєР° СЂРµРєРѕРјРµРЅРґР°С†РёР№", value=True)

# РџРѕР»Рµ РґР»СЏ Р·Р°РіСЂСѓР·РєРё С„Р°Р№Р»Р°
uploaded_file = st.file_uploader("Р—Р°РіСЂСѓР·РёС‚Рµ С„Р°Р№Р» СЂРµР·СЋРјРµ (PDF РёР»Рё DOCX)", type=["pdf", "docx"])

# РћР±СЂР°Р±РѕС‚РєР° Р·Р°РіСЂСѓР¶РµРЅРЅРѕРіРѕ С„Р°Р№Р»Р°
if uploaded_file is not None:
    st.success(f"Р¤Р°Р№Р» '{uploaded_file.name}' СѓСЃРїРµС€РЅРѕ Р·Р°РіСЂСѓР¶РµРЅ!")
    
    # РџСЂРѕРіСЂРµСЃСЃ-Р±Р°СЂ
    file_processing_progress_bar()

    file_text = ""

    if uploaded_file.name.endswith(".pdf"):
        file_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.name.endswith(".docx"):
        file_text = extract_text_from_docx(uploaded_file)

    if file_text:
        # РћРїСЂРµРґРµР»РµРЅРёРµ РєР°С‚РµРіРѕСЂРёРё СЂРµР·СЋРјРµ
        category = determine_resume_category(file_text)
        st.subheader("РљР°С‚РµРіРѕСЂРёСЏ СЂРµР·СЋРјРµ")
        st.write(f"РћРїСЂРµРґРµР»РµРЅР° РєР°С‚РµРіРѕСЂРёСЏ: **{category.capitalize()}**")

        # Р“РµРЅРµСЂР°С†РёСЏ СЂРµРєРѕРјРµРЅРґР°С†РёР№
        st.subheader("Р РµРєРѕРјРµРЅРґР°С†РёРё РїРѕ СѓР»СѓС‡С€РµРЅРёСЋ СЂРµР·СЋРјРµ")
        
        # РРЅС‚РµСЂР°РєС‚РёРІРЅР°СЏ РЅР°СЃС‚СЂРѕР№РєР° СЂРµРєРѕРјРµРЅРґР°С†РёР№
        recommendation_type = interactive_recommendations()
        if recommendation_type == "Р”Р»СЏ СЃС‚СѓРґРµРЅС‚РѕРІ":
            recommendations = generate_recommendations("student")
        elif recommendation_type == "Р”Р»СЏ РїСЂРѕС„РµСЃСЃРёРѕРЅР°Р»РѕРІ":
            recommendations = generate_recommendations("professional")
        else:
            recommendations = generate_recommendations("general")

        # Р’С‹РІРѕРґ СЂРµРєРѕРјРµРЅРґР°С†РёР№ СЃ С†РІРµС‚РѕРІРѕР№ РєРѕРґРёСЂРѕРІРєРѕР№
        for rec in recommendations:
            if color_coding:
                st.markdown(f"<p style='color: #4CAF50; font-size: 16px;'>{rec}</p>", unsafe_allow_html=True)
            else:
                st.markdown(f"<p style='font-size: 16px;'>{rec}</p>", unsafe_allow_html=True)

        # РљРЅРѕРїРєР° РґР»СЏ СЃРєР°С‡РёРІР°РЅРёСЏ PDF
        pdf_file = generate_pdf(recommendations)
        st.download_button("РЎРєР°С‡Р°С‚СЊ СЂРµРєРѕРјРµРЅРґР°С†РёРё (PDF)", pdf_file, file_name="recommendations.pdf", mime="application/pdf")

        # РљРЅРѕРїРєР° РґР»СЏ СЃРєР°С‡РёРІР°РЅРёСЏ DOCX
        docx_file = generate_docx(recommendations)
        st.download_button("РЎРєР°С‡Р°С‚СЊ СЂРµРєРѕРјРµРЅРґР°С†РёРё (DOCX)", docx_file, file_name="recommendations.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.error("РќРµ СѓРґР°Р»РѕСЃСЊ РёР·РІР»РµС‡СЊ С‚РµРєСЃС‚ РёР· С„Р°Р№Р»Р°. РџСЂРѕРІРµСЂСЊС‚Рµ РµРіРѕ СЃРѕРґРµСЂР¶РёРјРѕРµ.")